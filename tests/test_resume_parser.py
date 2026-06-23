from io import BytesIO
from pathlib import Path

from app.intake import resume_parser as rp


def test_extract_text_plaintext():
    data = "王小明\nPython 後端工程師".encode("utf-8")
    assert "Python" in rp.extract_text(data, "resume.txt")


def test_extract_text_unknown_ext_treated_as_text():
    data = "純文字履歷".encode("utf-8")
    assert rp.extract_text(data, "resume.unknown") == "純文字履歷"


def test_extract_text_docx():
    from docx import Document
    doc = Document()
    doc.add_paragraph("王小明")
    doc.add_paragraph("Python 後端工程師，3 年經驗")
    buf = BytesIO()
    doc.save(buf)
    text = rp.extract_text(buf.getvalue(), "resume.docx")
    assert "王小明" in text
    assert "後端工程師" in text


def test_extract_text_pdf():
    data = Path("tests/fixtures/sample_resume.pdf").read_bytes()
    text = rp.extract_text(data, "resume.pdf")
    assert "Resume" in text

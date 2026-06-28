"""投遞包 → Word（.docx）：用 python-docx 產出可在 Word 編輯的求職投遞包。

pkg 形狀（皆為選填）：
  {
    "job_title": str, "company": str,
    "resume": {"summary": str, "bullets": [str], "ats_keywords_hit": [str]},
    "cover_letter": {"subject": str, "body": str},
    "interview": {"technical_questions": [str], "behavioral_questions": [str], ...},
  }
中文字型走 w:eastAsia 提示（Windows 上 Word 以 Microsoft JhengHei 呈現；缺字型則自動代換）。
"""
from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

_INTERVIEW_SECTIONS = [
    ("技術題", "technical_questions"),
    ("行為題", "behavioral_questions"),
    ("歐洲特有題", "eu_specific_questions"),
    ("STAR 擬答", "sample_answers"),
    ("反向提問", "reverse_questions"),
    ("避雷提醒", "cautions"),
]


def _set_cjk_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), "Microsoft JhengHei")


def _bullets(doc: Document, items) -> None:
    if not isinstance(items, (list, tuple)):
        return
    for it in items:
        if str(it).strip():
            doc.add_paragraph(str(it), style="List Bullet")


def _section(pkg: dict, key: str) -> dict:
    """安全取出區塊：非 dict（字串/list/None）一律視為缺，不拋例外。"""
    v = pkg.get(key)
    return v if isinstance(v, dict) else {}


def build_docx(pkg: dict) -> bytes:
    if not isinstance(pkg, dict):
        pkg = {}
    doc = Document()
    _set_cjk_font(doc)

    doc.add_heading(str(pkg.get("job_title") or "求職投遞包"), level=0)
    if pkg.get("company"):
        doc.add_paragraph(str(pkg["company"]))

    resume = _section(pkg, "resume")
    if resume:
        doc.add_heading("客製履歷", level=1)
        if resume.get("summary"):
            doc.add_paragraph(str(resume["summary"]))
        _bullets(doc, resume.get("bullets"))
        hits = resume.get("ats_keywords_hit")
        if isinstance(hits, (list, tuple)):
            cleaned = [str(h) for h in hits if str(h).strip()]
            if cleaned:
                doc.add_paragraph("ATS 命中關鍵字：" + "、".join(cleaned))

    cover = _section(pkg, "cover_letter")
    if cover:
        doc.add_heading("求職信", level=1)
        if cover.get("subject"):
            p = doc.add_paragraph()
            p.add_run("主旨：" + str(cover["subject"])).bold = True
        for line in str(cover.get("body") or "").split("\n"):
            doc.add_paragraph(line)

    interview = _section(pkg, "interview")
    if interview:
        doc.add_heading("面試準備", level=1)
        for label, key in _INTERVIEW_SECTIONS:
            items = interview.get(key)
            if items:
                doc.add_heading(label, level=2)
                _bullets(doc, items)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
